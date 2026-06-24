"""
make_bad_word.py — build a deliberately flawed .docx to validate tools/check_word.py.
Usage: python examples/make_bad_word.py [out.docx]   (default ./_test_flawed_doc.docx)
"""
import sys
import os
from docx import Document
from docx.shared import Inches
from PIL import Image

HERE = os.path.dirname(os.path.abspath(__file__))


def build(out_path):
    doc = Document()

    # --- clean control: proper H1 -> H2 sequence, normal text (NO finding) ---
    doc.add_heading("Real Title", level=1)
    doc.add_heading("A Proper Subsection", level=2)
    doc.add_paragraph("This is ordinary, clean body text. Nothing wrong here.")

    # --- DOC-HEADING-SKIP: H1 -> H4 (skips 2,3) ---
    doc.add_heading("Jumps Too Deep", level=4)

    # --- DOC-EMPTY-HEADING: a heading with no text ---
    doc.add_heading("", level=2)

    # --- DOC-MARKDOWN: literal markdown emitted into Word ---
    doc.add_paragraph("**bold left as literal markdown**")
    doc.add_paragraph("## A heading written as literal hashes")
    doc.add_paragraph("- a dash bullet that should have been a real list")
    doc.add_paragraph("See the [docs](https://example.com) link in raw markdown")

    # --- DOC-PLACEHOLDER: boilerplate + unresolved token ---
    doc.add_paragraph("TODO: write this section before sending")
    doc.add_paragraph("Dear {{client_name}}, thank you for ${amount}.")
    doc.add_paragraph("[INSERT DATE HERE]")

    # --- DOC-FIELD-ERR: broken field + mojibake ---
    doc.add_paragraph("As shown in Error! Reference source not found. above.")
    doc.add_paragraph("Encoding broke here: caf� and na�ve")

    # --- DOC-EMPTY-TABLE: 2x2 all-empty table ---
    doc.add_table(rows=2, cols=2)

    # --- DOC-TABLE-RAGGED: a table whose 2nd row has one fewer cell ---
    t = doc.add_table(rows=2, cols=3)
    t.cell(0, 0).text = "a"; t.cell(0, 1).text = "b"; t.cell(0, 2).text = "c"
    t.cell(1, 0).text = "1"; t.cell(1, 1).text = "2"; t.cell(1, 2).text = "3"
    # remove the last cell of row 2 -> ragged
    last_tc = t.rows[1].cells[2]._tc
    last_tc.getparent().remove(last_tc)

    # --- DOC-IMG-DISTORT: 4:1 native image placed as a square ---
    img = os.path.join(HERE, "_sample_4x1.png")
    if not os.path.exists(img):
        Image.new("RGB", (400, 100), (40, 90, 200)).save(img)
    doc.add_picture(img, width=Inches(2), height=Inches(2))  # squished to 1:1

    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    out = sys.argv[1] if len(sys.argv) > 1 else os.path.join(HERE, "_test_flawed_doc.docx")
    print("wrote", build(out))
