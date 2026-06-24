"""
check_word.py — objective-defect checker for .docx documents.

Same scope philosophy as Office File Inspector: flag ONLY objective,
unambiguous mistakes a strictly MORE capable model would also never want.
Two tests per check: (1) OBJECTIVE  (2) NO-SHACKLE.

Word reflows text, so most "layout" defects don't exist here — the real LLM
defects are GENERATION ARTIFACTS (markdown emitted into a Word doc, broken
fields, leftover placeholders) and a few STRUCTURAL ones (empty/skipped headings,
ragged tables, stretched images).

NOT CHECKED (style / shackle, or needs the user's baseline = process-discipline):
  font/spacing/quote style, direct-formatting vs styles, "color looks off",
  paragraph density, line length.

Checks:
  DOC-FIELD-ERR   broken field result ("Error! Reference source not found"),
                  mojibake replacement char                                  ERROR
  DOC-TABLE-RAGGED rows of a table have differing cell counts                ERROR
  DOC-MARKDOWN    literal markdown left in the text (**), #, -, |, [](...)   WARN
  DOC-PLACEHOLDER leftover placeholder / unresolved {{token}} / ${var}       WARN
  DOC-EMPTY-HEADING a Heading-styled paragraph with no text                  WARN
  DOC-HEADING-SKIP outline jumps more than one level (H1 -> H3)              WARN
  DOC-EMPTY-TABLE  a table whose cells are all empty                         WARN
  DOC-IMG-DISTORT  inline image stretched off its native aspect ratio        WARN

Exit: 1 if any ERROR (or any WARN with --strict). 0 otherwise. 2 on usage error.

Usage:
    python tools/check_word.py doc.docx [--json report.json] [--strict] [--quiet]
"""
import argparse
import json
import os
import re
import sys

from docx import Document
from docx.enum.shape import WD_INLINE_SHAPE

MD_PATTERNS = [
    (re.compile(r"\*\*[^*]+\*\*"), "literal **bold** markdown"),
    (re.compile(r"__[^_]+__"), "literal __bold__ markdown"),
    (re.compile(r"^\s{0,3}#{1,6}\s+\S"), "literal #-heading markdown"),
    (re.compile(r"^\s{0,3}[-*+]\s+\S"), "literal dash/asterisk bullet"),
    (re.compile(r"^\s*\|.+\|.+\|\s*$"), "literal pipe-table markdown"),
    (re.compile(r"\[[^\]]+\]\([^)]+\)"), "literal [link](url) markdown"),
]
PLACEHOLDER_TEXT = re.compile(
    r"click to add|lorem ipsum|your text here|insert .* here|sample text|"
    r"todo:|tbd\b|xxx+|fixme|\[(insert|name|date|company)[^\]]*\]",
    re.I,
)
TOKEN_RE = re.compile(r"\{\{[^}]+\}\}|\$\{[^}]+\}|<<[^>]+>>|%[A-Z_]{3,}%")
FIELD_ERR_RE = re.compile(
    r"Error!\s+(Reference source not found|Bookmark not defined|"
    r"Hyperlink reference not valid|No (table|figure) of)", re.I)


def heading_level(style_name):
    if not style_name:
        return None
    if style_name == "Title":
        return 0
    m = re.match(r"Heading (\d+)", style_name)
    return int(m.group(1)) if m else None


class Linter:
    def __init__(self):
        self.findings = []

    def add(self, sev, cid, where, msg):
        self.findings.append({"id": cid, "severity": sev, "where": where, "message": msg})

    def run(self, path):
        doc = Document(path)
        self.check_paragraphs(doc)
        self.check_tables(doc)
        self.check_images(doc)
        return {
            "file": os.path.abspath(path),
            "summary": self.summary(),
            "findings": sorted(self.findings, key=lambda f: (f["id"], f["where"])),
        }

    def summary(self):
        s = {"ERROR": 0, "WARN": 0, "INFO": 0}
        for f in self.findings:
            s[f["severity"]] += 1
        return s

    def check_text(self, text, where):
        if not text or not text.strip():
            return
        if FIELD_ERR_RE.search(text) or "�" in text:
            reason = "broken field reference" if FIELD_ERR_RE.search(text) else "mojibake / replacement char"
            self.add("ERROR", "DOC-FIELD-ERR", where, f"{reason}: {text.strip()[:50]!r}")
        for pat, why in MD_PATTERNS:
            if pat.search(text):
                self.add("WARN", "DOC-MARKDOWN", where, f"{why}: {text.strip()[:50]!r}")
                break
        if PLACEHOLDER_TEXT.search(text) or TOKEN_RE.search(text):
            self.add("WARN", "DOC-PLACEHOLDER", where, f"placeholder/unresolved token: {text.strip()[:50]!r}")

    def check_paragraphs(self, doc):
        prev_level = None
        for i, p in enumerate(doc.paragraphs):
            where = f"para {i+1}"
            lvl = heading_level(p.style.name if p.style else None)
            if lvl is not None and lvl >= 1:
                if not p.text.strip():
                    self.add("WARN", "DOC-EMPTY-HEADING", where,
                             f"empty '{p.style.name}' heading.")
                if prev_level is not None and lvl > prev_level + 1:
                    self.add("WARN", "DOC-HEADING-SKIP", where,
                             f"outline jumps H{prev_level} -> H{lvl} (level skipped).")
                prev_level = lvl
            self.check_text(p.text, where)

    def check_tables(self, doc):
        for ti, t in enumerate(doc.tables):
            counts = [len(r.cells) for r in t.rows]
            if len(set(counts)) > 1:
                self.add("ERROR", "DOC-TABLE-RAGGED", f"table {ti+1}",
                         f"rows have differing cell counts {counts}.")
            nonempty = False
            for ri, r in enumerate(t.rows):
                for ci, c in enumerate(r.cells):
                    if c.text.strip():
                        nonempty = True
                    self.check_text(c.text, f"table {ti+1} cell({ri+1},{ci+1})")
            if not nonempty and counts:
                self.add("WARN", "DOC-EMPTY-TABLE", f"table {ti+1}", "table is entirely empty.")

    def check_images(self, doc):
        for si, shape in enumerate(doc.inline_shapes):
            try:
                if shape.type != WD_INLINE_SHAPE.PICTURE:
                    continue
                dw, dh = shape.width, shape.height
                if not dw or not dh:
                    continue
                blip = shape._inline.graphic.graphicData.pic.blipFill.blip
                rId = blip.embed
                part = doc.part.related_parts[rId]
                img = part.image
                nw, nh = img.px_width, img.px_height
            except Exception:
                continue
            if not nw or not nh:
                continue
            disp = dw / dh
            native = nw / nh
            if native == 0:
                continue
            dev = abs(disp / native - 1.0)
            if dev > 0.04:
                self.add("WARN", "DOC-IMG-DISTORT", f"inline image {si+1}",
                         f"stretched: displayed ratio {disp:.2f} vs native {native:.2f} "
                         f"({dev*100:.0f}% off).")


SEV_ORDER = {"ERROR": 0, "WARN": 1, "INFO": 2}
SEV_MARK = {"ERROR": "x", "WARN": "!", "INFO": "i"}


def print_human(report):
    s = report["summary"]
    print(f"\nDOCX LINT  {os.path.basename(report['file'])}")
    print(f"  {s['ERROR']} ERROR   {s['WARN']} WARN   {s['INFO']} INFO\n")
    if not report["findings"]:
        print("  clean -- no objective defects found.\n")
        return
    for f in sorted(report["findings"], key=lambda x: (SEV_ORDER[x["severity"]], x["id"], x["where"])):
        print(f"  [{SEV_MARK[f['severity']]}] {f['id']} [{f['where']}]: {f['message']}")
    print()


def main():
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
        sys.stderr.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass
    ap = argparse.ArgumentParser()
    ap.add_argument("docx")
    ap.add_argument("--json", default=None)
    ap.add_argument("--strict", action="store_true")
    ap.add_argument("--quiet", action="store_true")
    args = ap.parse_args()

    if not os.path.isfile(args.docx):
        print("ERROR: no such file:", args.docx, file=sys.stderr)
        sys.exit(2)

    report = Linter().run(args.docx)
    if args.json:
        with open(args.json, "w", encoding="utf-8") as fh:
            json.dump(report, fh, ensure_ascii=False, indent=2)
        print("wrote", args.json)
    if not args.quiet:
        print_human(report)

    s = report["summary"]
    sys.exit(1 if s["ERROR"] > 0 or (args.strict and s["WARN"] > 0) else 0)


if __name__ == "__main__":
    main()
